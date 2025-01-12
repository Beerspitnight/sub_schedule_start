# aesop_parser.py

from typing import List, Dict, Any
import logging
from docx import Document
import re
from fuzzywuzzy import fuzz, process
from io import BytesIO
from googleapiclient.http import MediaIoBaseDownload

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Google Drive file IDs
TEACHERLIST_ID = "14PEFZHKCTyvtc8O-xSThvPvbPVLSd4BaGadk7UW6tGw"
SUBLIST_ID = "10khvv6YZacm7lkNjq1gsJky3BpHqguR5C6dMjDr1Qjg"
MASTER_SCHEDULE_ID = "12XNbaa4AvahxYxR7D6Qa6DfrEeZPgrSiTTo5V9uFTsg"
DAILY_COVERAGE_ID = "1vFsNkhIVqaS72JD5PV1nOSlYYME4XTpJBBai8d65S84"

def parse_aesop_report_names(service, folder_id: str, teacherlist: List[str], sublist: List[str], threshold: int = 80) -> Dict[str, List[str]]:
    """
    Parses the latest Aesop Daily Report from Google Drive folder to find matches for teacher and substitute names.
    
    Args:
        service: Google Drive service instance
        folder_id: ID of the folder containing the Aesop report
        teacherlist: List of teacher names to search for
        sublist: List of substitute names to search for
        threshold: Minimum fuzzy match score to consider a match valid
        
    Returns:
        Dictionary with matched absent teachers and assigned substitutes
    """
    try:
        # Get the latest Aesop report file ID using drive_utilities
        from drive_utilities import get_latest_docx_in_folder
        
        file_id = get_latest_docx_in_folder(folder_id, service)
        if not file_id:
            raise FileNotFoundError("No Aesop report found in the specified folder")
        
        # Download and read the file content
        request = service.files().get_media(fileId=file_id)
        file_content = BytesIO()
        downloader = MediaIoBaseDownload(file_content, request)
        
        done = False
        while done is False:
            status, done = downloader.next_chunk()
            
        file_content.seek(0)
        
        # Load the document
        doc = Document(file_content)
        
        # Extract all text from paragraphs and tables
        text_content = []
        
        # Get paragraph text and table text
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
                
        for table in doc.tables:
            for row in table.rows:
                text_content.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
                
        # Join all text into single string
        document_text = " ".join(text_content)
        
        # Initialize results
        found_matches = {
            "absent_teachers": [],
            "assigned_subs": []
        }
        
        # Search for teachers
        for teacher in teacherlist:
            matches = process.extract(teacher, text_content, scorer=fuzz.token_sort_ratio)
            good_matches = [match for match in matches if match[1] >= threshold]
            if good_matches:
                found_matches["absent_teachers"].append(teacher)
                logger.debug(f"Found teacher match: {teacher}")
                
        # Search for substitutes
        for sub in sublist:
            matches = process.extract(sub, text_content, scorer=fuzz.token_sort_ratio)
            good_matches = [match for match in matches if match[1] >= threshold]
            if good_matches:
                found_matches["assigned_subs"].append(sub)
                logger.debug(f"Found substitute match: {sub}")
                
        logger.info(f"Found {len(found_matches['absent_teachers'])} teachers and {len(found_matches['assigned_subs'])} substitutes")
        return found_matches
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise

if __name__ == "__main__":
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    import json
    
    # Load credentials from local JSON file
    with open('credentials.json', 'r') as f:
        creds_data = json.load(f)
    credentials = Credentials.from_authorized_user_info(creds_data)
    
    # Build the Drive service
    drive_service = build('drive', 'v3', credentials=credentials)
    
    try:
        # Get your lists from their respective sources
        # (You'll need to implement the actual data fetching)
        teacherlist = []  # Load from TEACHERLIST_ID
        sublist = []      # Load from SUBLIST_ID
        
        # Parse the document
        matches = parse_aesop_report_names(
            service=drive_service,
            folder_id=""1RJzDwcEluGkqglmX-e3cg09Ai35kBSQG"",  # Replace with actual folder ID containing Aesop reports
            teacherlist=teacherlist,
            sublist=sublist
        )
        
        # Print results
        print("\nAbsent Teachers:", matches["absent_teachers"])
        print("Assigned Substitutes:", matches["assigned_subs"])
        
    except Exception as e:
        logger.error(f"Failed to process report: {e}")